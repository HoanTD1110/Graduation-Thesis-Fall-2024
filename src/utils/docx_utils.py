from docx import *
from docx.text.paragraph import Paragraph
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pandas as pd
import io
import csv


# This function extracts the tables and paragraphs from the document object
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# This function extracts the table from the document object as a dataframe
def read_docx_tables(document, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """

    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    #    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in document.tables]
    else:
        try:
            return read_docx_tab(document.tables[tab_id], **kwargs)
        except IndexError:
            print(
                "Error: specified [tab_id]: {}  does not exist.".format(tab_id))
            raise


def extract_docx(filename):
    document = Document(filename)

    combined_df = pd.DataFrame(columns=["para_text", "table_id", "style"])
    table_mod = pd.DataFrame(columns=["string_value", "table_id"])
    # The table_list is a list consisting of all the tables in the document
    table_list = []
    i = 0

    for block in iter_block_items(document):
        isappend = None
        tabid = None
        appendtxt = None
        style = None
        if "text" in str(block):
            isappend = False

            runboldtext = ""
            for run in block.runs:
                if run.bold:
                    runboldtext = runboldtext + run.text

            style = str(block.style.name)
            appendtxt = str(block.text)
            appendtxt = appendtxt.replace("\n", "")
            appendtxt = appendtxt.replace("\r", "")
            tabid = -1
            isappend = True

        elif "table" in str(block):
            isappend = True
            style = "Novalue"
            appendtxt = str(block)
            tabid = i
            dfs = read_docx_tables(document, tab_id=i)
            dftemp = pd.DataFrame(
                {"para_text": [appendtxt], "table_id": [i], "style": [style]}
            )
            table_mod = pd.concat([table_mod, dftemp], sort=False)
            table_list.append(dfs)
            i = i + 1
        if isappend:
            dftemp = pd.DataFrame(
                {"para_text": [appendtxt], "table_id": [
                    tabid], "style": [style]}
            )
            combined_df = pd.concat([combined_df, dftemp], sort=False)

    combined_df = combined_df.reset_index(drop=True)
    return combined_df
